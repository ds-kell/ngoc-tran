from pathlib import Path
from collections import Counter
from docx import Document

root = Path(r"D:\NgocTran\phase1")
work = root / "_work_knk_2022"
work.mkdir(exist_ok=True)

files = [
    root / "2025_11_10 BTR1_Chuong 2_draft2.docx",
    root / "bao cao chuyen de" / "2025-4-29. IE. Bao cao KKQG KNK nam 2021, 2022_nang luong.docx",
    root / "bao cao chuyen de" / "Báo cáo KK KNK 2021-2022_IPPU_ver0811.docx",
    work / "CĐ1_4.1 Độ không chắc chắn _ Năng lượng.docx",
    work / "CĐ4_4.4 Độ không chắc chắn _ IPPU.docx",
]

for path in files:
    doc = Document(str(path))
    lines = [f"SOURCE: {path}", f"PARAGRAPHS={len(doc.paragraphs)} TABLES={len(doc.tables)} SECTIONS={len(doc.sections)}", ""]
    styles = Counter()
    fonts = Counter()
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt:
            lines.append(f"P{i:04d} [{p.style.name if p.style else ''}] {txt}")
        styles[p.style.name if p.style else ""] += 1
        for r in p.runs:
            if r.text.strip():
                fonts[(r.font.name or "", round(r.font.size.pt, 1) if r.font.size else None, bool(r.bold), bool(r.italic))] += len(r.text)
    for ti, table in enumerate(doc.tables):
        lines.append(f"\n=== TABLE {ti+1} ({len(table.rows)}x{len(table.columns)}) ===")
        for ri, row in enumerate(table.rows):
            vals = [" / ".join(p.text.strip() for p in c.paragraphs if p.text.strip()) for c in row.cells]
            lines.append(f"R{ri:03d}\t" + "\t".join(vals))
    lines.append("\n=== STYLE COUNTS ===")
    lines.extend(f"{k}\t{v}" for k,v in styles.most_common())
    lines.append("\n=== DIRECT FONT COUNTS (characters) ===")
    lines.extend(f"{k}\t{v}" for k,v in fonts.most_common(30))
    out = work / (path.stem + ".txt")
    out.write_text("\n".join(lines), encoding="utf-8")
    print(path.stem[:8].encode("ascii", "replace").decode(), len(lines))

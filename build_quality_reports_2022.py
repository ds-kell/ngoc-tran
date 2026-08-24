from pathlib import Path
import re
import shutil

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import update_knk_docs as base

ROOT = Path(r"D:\NgocTran\phase1")
TARGET = ROOT / "Đảm bảo chất lượng 2022"
MD_DIR = TARGET / "md"
WORK = ROOT / "_work_knk_2022"
BACKUP = TARGET / "_backup_truoc_khi_lam_lai_dung_yeu_cau"

PAIRS = {
    "BC19_3.1.1": "BC19_3.1.1. Đảm bảo chất lượng số liệu hoạt động trong lĩnh vực năng lượng.docx",
    "BC20_3.1.2": "BC20_3.1.2. Đảm bảo chất lượng hệ số phát thải trong lĩnh vực năng lượng.docx",
    "BC21_3.1.3": "BC21_3.1.3. Đảm bảo chất lượng kết quả KKKNK trong lĩnh vực năng lượng.docx",
    "BC28_3.4.1": "BC28_3.4.1. Đảm bảo chất lượng số liệu hoạt động trong lĩnh vực IPPU.docx",
    "BC29_3.4.2": "BC29_3.4.2. Đảm bảo chất lượng hệ số phát thải trong lĩnh vực IPPU.docx",
    "BC30_3.4.3": "BC30_3.4.3. Đảm bảo chất lượng kết quả kiểm kê KNK trong lĩnh vực IPPU.docx",
}


def add_field(paragraph, instruction):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Nhấn F9 để cập nhật"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])


def preprocess_markdown(text):
    text = text.replace("Hà Nội, 2026", "Hà Nội, 2025")
    # Use real Word fields instead of the Markdown's static lists.
    text = re.sub(r"## MỤC LỤC\n.*?(?=\n## DANH MỤC BẢNG)", "## MỤC LỤC\n\n[[TOC]]\n", text, flags=re.S)
    text = re.sub(r"## DANH MỤC BẢNG\n.*?(?=\n## DANH MỤC TỪ VIẾT TẮT)", "## DANH MỤC BẢNG\n\n[[LOT]]\n", text, flags=re.S)
    return text


def set_run_font(run, size=14, bold=None, italic=None):
    base.set_font(run, "Times New Roman", size=size, bold=bold, italic=italic)


def style_document(doc):
    # Geometry follows the converted uncertainty templates, while normalizing
    # the body to the dominant template typography (TNR 14).
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size in (("Heading 1", 16), ("Heading 2", 15), ("Heading 3", 14)):
        if name in doc.styles:
            st = doc.styles[name]
            st.font.name = "Times New Roman"
            st.font.size = Pt(size)
            st.font.bold = True
            st.font.color.rgb = RGBColor(0, 0, 0)
            st.paragraph_format.space_before = Pt(12)
            st.paragraph_format.space_after = Pt(6)
            st.paragraph_format.keep_with_next = True
    for pi, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        # Functional TOC/list-of-tables fields.
        if txt == "[[TOC]]":
            p.clear()
            add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
        elif txt == "[[LOT]]":
            p.clear()
            add_field(p, 'TOC \\h \\z \\t "Caption,1"')
        # Table captions become a real caption style for the list of tables.
        if re.match(r"^Bảng\s+\d", txt, re.I):
            try:
                p.style = doc.styles["Caption"]
            except KeyError:
                pass
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            for r in p.runs:
                set_run_font(r, 13, bold=True)
        # Major divisions begin on a new page, mirroring the reference reports.
        # MỤC LỤC already follows the explicit cover-page break in Markdown.
        if txt in {"DANH MỤC BẢNG", "DANH MỤC TỪ VIẾT TẮT", "MỞ ĐẦU", "KẾT LUẬN", "TÀI LIỆU THAM KHẢO"} or txt.startswith("CHƯƠNG "):
            p.paragraph_format.page_break_before = True
        # Cover paragraphs are centered and slightly more open.
        if pi < 9:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            if r.text.strip() and r.font.size is None:
                set_run_font(r, 14)
    for table in doc.tables:
        table.autofit = True
        for ri, row in enumerate(table.rows):
            trpr = row._tr.get_or_add_trPr()
            if ri == 0 and trpr.find(qn("w:tblHeader")) is None:
                trpr.append(OxmlElement("w:tblHeader"))
            for cell in row.cells:
                base.set_cell_margins(cell, top=90, start=90, bottom=90, end=90)
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_after = Pt(0)
                    for r in p.runs:
                        set_run_font(r, 10.5, bold=(True if ri == 0 else r.bold))
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main():
    BACKUP.mkdir(exist_ok=True)
    for old in TARGET.glob("BC*.docx"):
        if not (BACKUP / old.name).exists():
            shutil.copy2(old, BACKUP / old.name)
    energy_template = WORK / "CĐ1_4.1 Độ không chắc chắn _ Năng lượng.docx"
    ippu_template = WORK / "CĐ4_4.4 Độ không chắc chắn _ IPPU.docx"
    for code, name in PAIRS.items():
        template = energy_template if code.startswith(("BC19", "BC20", "BC21")) else ippu_template
        doc = Document(str(template))
        base.clear_body(doc)
        base.configure_document(doc)
        text = preprocess_markdown((MD_DIR / f"{code}.md").read_text(encoding="utf-8"))
        base.add_markdown(doc, text)
        style_document(doc)
        out = TARGET / name
        doc.save(str(out))
        print(code)


if __name__ == "__main__":
    main()

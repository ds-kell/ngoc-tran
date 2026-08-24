from pathlib import Path
import re
from docx import Document

root = Path(r"D:\NgocTran\phase1\Chuyên đề kiểm kê KNKQG 2023")
out = root / "Đảm bảo chất lượng 2023"

def norm(s):
    s = re.sub(r"[*`>#|]", "", s)
    s = re.sub(r"^\s*[-+]\s+", "", s)
    s = re.sub(r"^\s*\d+\.\s+", "", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s

for md in sorted(root.glob("*.md")):
    prefix = md.stem
    docx = next(out.glob(prefix + "*.docx"))
    doc = Document(str(docx))
    body = "\n".join([p.text for p in doc.paragraphs] + [c.text for t in doc.tables for row in t.rows for c in row.cells])
    body_n = norm(body)
    meaningful = []
    md_tables = 0
    in_table = False
    for line in md.read_text(encoding="utf-8").splitlines():
        s = line.strip()
        if s.startswith("|"):
            if not in_table:
                md_tables += 1
                in_table = True
        else:
            in_table = False
        n = norm(s)
        if not n or n == "---" or re.fullmatch(r":?-{3,}:?(\s+:?-{3,}:?)*", n):
            continue
        meaningful.append(n)
    missing = [x for x in meaningful if x not in body_n and not re.fullmatch(r"[-: ]+", x)]
    print(f"{prefix}: paragraphs={len(doc.paragraphs)}, tables={len(doc.tables)}/{md_tables}, missing={len(missing)}")
    for x in missing[:3]:
        print("  MISSING:", x[:160])

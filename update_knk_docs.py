from pathlib import Path
import re
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"D:\NgocTran\phase1\Chuyên đề kiểm kê KNKQG 2023")
SRC = ROOT / "Đảm bảo chất lượng 2022"
OUT = ROOT / "Đảm bảo chất lượng 2023"

PAIRS = {
    "BC19_3.1.1": "BC19_3.1.1. Đảm bảo chất lượng số liệu hoạt động trong lĩnh vực năng lượng.docx",
    "BC20_3.1.2": "BC20_3.1.2. Đảm bảo chất lượng hệ số phát thải trong lĩnh vực năng lượng.docx",
    "BC21_3.1.3": "BC21_3.1.3. Đảm bảo chất lượng kết quả KKKNK trong lĩnh vực năng lượng.docx",
    "BC28_3.4.1": "BC28_3.4.1. Đảm bảo chất lượng số liệu hoạt động trong lĩnh vực IPPU.docx",
    "BC29_3.4.2": "BC29_3.4.2. Đảm bảo chất lượng hệ số phát thải trong lĩnh vực IPPU.docx",
    "BC30_3.4.3": "BC30_3.4.3. Đảm bảo chất lượng kết quả kiểm kê KNK trong lĩnh vực IPPU.docx",
}


def set_font(run, name="Times New Roman", size=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.rFonts
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rf.set(qn("w:" + key), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def clear_body(doc):
    body = doc._element.body
    sect = body.sectPr
    for child in list(body):
        if child is not sect:
            body.remove(child)


def add_inline(paragraph, text, default_bold=False, default_italic=False):
    # Minimal Markdown inline handling: bold, italic, bold-italic and code.
    token_re = re.compile(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*[^*]+?\*|`[^`]+?`)")
    pos = 0
    for m in token_re.finditer(text):
        if m.start() > pos:
            set_font(paragraph.add_run(text[pos:m.start()]), bold=default_bold, italic=default_italic)
        tok = m.group(0)
        bold, italic = default_bold, default_italic
        if tok.startswith("***"):
            val, bold, italic = tok[3:-3], True, True
        elif tok.startswith("**"):
            val, bold = tok[2:-2], True
        elif tok.startswith("*"):
            val, italic = tok[1:-1], True
        else:
            val = tok[1:-1]
        set_font(paragraph.add_run(val), bold=bold, italic=italic)
        pos = m.end()
    if pos < len(text):
        set_font(paragraph.add_run(text[pos:]), bold=default_bold, italic=default_italic)


def shade_cell(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    mar = tcpr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcpr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = mar.find(qn("w:" + side))
        if el is None:
            el = OxmlElement("w:" + side)
            mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def parse_table(lines, i):
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def configure_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
    if "Normal" not in doc.styles:
        doc.styles.add_style("Normal", WD_STYLE_TYPE.PARAGRAPH)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, color in (("Heading 1", 15, "17365D"), ("Heading 2", 14, "1F4E79"), ("Heading 3", 13, "365F91")):
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        st = doc.styles[style_name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(5)
        st.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Bullet 2", "List Bullet 3", "List Number"):
        if style_name not in doc.styles:
            st = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            st.base_style = normal
            st.paragraph_format.left_indent = Cm(0.7 if style_name in ("List Bullet", "List Number") else 1.4)
            st.paragraph_format.first_line_indent = Cm(-0.35)
    if "Table Grid" not in doc.styles:
        doc.styles.add_style("Table Grid", WD_STYLE_TYPE.TABLE)


def set_table_borders(table, color="8FA8BC", size="6"):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn("w:" + edge))
        if el is None:
            el = OxmlElement("w:" + edge)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def add_markdown(doc, text):
    lines = text.replace("\r\n", "\n").split("\n")
    i = 0
    h1_count = 0
    on_cover = True
    while i < len(lines):
        raw = lines[i]
        s = raw.strip()
        if not s:
            i += 1
            continue
        if s == "---":
            if on_cover:
                doc.add_page_break()
                on_cover = False
            i += 1
            continue
        if s.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = parse_table(lines, i)
            if not rows:
                continue
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            set_table_borders(table)
            table.autofit = True
            for ri, row in enumerate(rows):
                for ci in range(cols):
                    cell = table.cell(ri, ci)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_margins(cell)
                    val = row[ci] if ci < len(row) else ""
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    if ci > 0 and re.fullmatch(r"[+−–—×\d\s.,%()*/]+", re.sub(r"\*", "", val)):
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_inline(p, val, default_bold=(ri == 0))
                    for run in p.runs:
                        set_font(run, size=10.5, bold=(True if ri == 0 else run.bold))
                    if ri == 0:
                        shade_cell(cell, "D9EAF7")
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", s)
        if m:
            level, val = len(m.group(1)), m.group(2)
            if level == 1:
                h1_count += 1
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if on_cover else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(10)
                add_inline(p, val, default_bold=True)
                for r in p.runs:
                    set_font(r, size=(15 if h1_count == 1 else 17), bold=True)
            else:
                p = doc.add_paragraph(style=f"Heading {min(level - 1, 3)}")
                add_inline(p, val)
            i += 1
            continue
        lm = re.match(r"^(\s*)[-+]\s+(.*)$", raw)
        if lm:
            depth = min(len(lm.group(1)) // 2, 2)
            style = "List Bullet" if depth == 0 else f"List Bullet {depth + 1}"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(0.7 * depth)
            add_inline(p, lm.group(2))
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue
        nm = re.match(r"^\s*(\d+)\.\s+(.*)$", raw)
        if nm:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, nm.group(2))
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue
        if s.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.right_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, s[1:].strip(), default_italic=True)
            i += 1
            continue
        # Join consecutive prose lines; Markdown source generally uses one line per paragraph.
        p = doc.add_paragraph()
        if on_cover:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
        add_inline(p, s)
        i += 1


def main():
    OUT.mkdir(exist_ok=True)
    for stem, docx_name in PAIRS.items():
        md = ROOT / f"{stem}.md"
        template = SRC / docx_name
        out = OUT / docx_name
        doc = Document(str(template))
        clear_body(doc)
        configure_document(doc)
        add_markdown(doc, md.read_text(encoding="utf-8"))
        # Ask Word to refresh fields (page numbers/TOC, if any remain in headers/footers).
        settings = doc.settings._element
        update = settings.find(qn("w:updateFields"))
        if update is None:
            update = OxmlElement("w:updateFields")
            settings.append(update)
        update.set(qn("w:val"), "true")
        doc.save(str(out))
        print(stem)


if __name__ == "__main__":
    main()

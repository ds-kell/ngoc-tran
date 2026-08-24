from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"D:\NgocTran\phase1\Đảm bảo chất lượng 2022")

def font(run, size=14):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.rFonts
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    for k in ("ascii", "hAnsi", "eastAsia", "cs"):
        rf.set(qn("w:" + k), "Times New Roman")

def add_page_field(paragraph):
    run = paragraph.add_run()
    font(run, 12)
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])

for path in sorted(ROOT.glob("BC*.docx")):
    doc = Document(str(path))
    # Remove the former mid-page year line from the body.
    for p in doc.paragraphs[:15]:
        if p.text.strip() in {"Hà Nội, 2025", "Hà Nội, 2026"}:
            p._element.getparent().remove(p._element)
            break
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.footer_distance = Cm(0.7)
    footer = section.first_page_footer
    for child in list(footer._element):
        footer._element.remove(child)
    p_year = footer.add_paragraph()
    p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_year.paragraph_format.space_after = Pt(4)
    r = p_year.add_run("Hà Nội, 2026")
    font(r, 14)
    p_num = footer.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_num.paragraph_format.space_before = Pt(0)
    p_num.paragraph_format.space_after = Pt(0)
    add_page_field(p_num)
    # Footer needs a final paragraph element to remain valid in Word.
    if not footer.paragraphs:
        footer.add_paragraph()
    doc.save(str(path))
    print(path.name[:4])
